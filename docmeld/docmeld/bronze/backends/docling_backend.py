"""Docling backend for document element extraction (PDF, DOCX, PPTX)."""

from __future__ import annotations

import base64
from pathlib import Path
from typing import Any


class DoclingBackend:
    """Extract elements using Docling's DocumentConverter + python-docx for OOXML.

    Docling handles body content (text, tables, titles, images, lists).
    python-docx handles header/footer/footnote text + page break detection
    from the OOXML ZIP structure.
    """

    def extract_elements(self, doc_path: str, output_dir: str) -> list[dict[str, Any]]:
        """Extract all elements from a document.

        Args:
            doc_path: Path to the document file (.pdf, .docx).
            output_dir: Directory for auxiliary outputs.

        Returns:
            List of element dicts in DocMeld format.
        """
        try:
            from docling.document_converter import DocumentConverter
        except ImportError as e:
            _msg = "Docling is not installed. Install with: pip install docmeld[docling]"
            raise ImportError(_msg) from e

        converter = DocumentConverter()
        result = converter.convert(doc_path)
        doc = result.document

        elements: list[dict[str, Any]] = []
        page_breaks: set[int] = set()
        is_docx = Path(doc_path).suffix.lower() == ".docx"

        # Extract headers/footers from OOXML for .docx
        hf_elements: list[dict[str, Any]] = []
        if is_docx:
            hf_elements = self._extract_ooxml_headers_footers(doc_path)
            page_breaks = self._detect_ooxml_page_breaks(doc_path)

        # Assign page numbers based on page break positions
        current_page = 1
        body_item_index = 0

        for item, _level in doc.iterate_items():
            item_type = type(item).__name__
            body_item_index += 1

            # Check if we've crossed a page break
            if page_breaks and body_item_index in page_breaks:
                current_page += 1

            page_no = self._get_page_no(item) or current_page

            if item_type in ("SectionHeaderItem", "TitleItem"):
                level = getattr(item, "level", 1)
                elements.append(
                    {
                        "type": "title",
                        "level": max(0, level - 1),
                        "content": item.text,
                        "page_no": page_no,
                    }
                )

            elif item_type == "TextItem":
                if item.text.strip():
                    elements.append(
                        {
                            "type": "text",
                            "content": item.text.strip(),
                            "page_no": page_no,
                        }
                    )

            elif item_type == "ListItem":
                if item.text.strip():
                    elements.append(
                        {
                            "type": "text",
                            "content": f"- {item.text.strip()}",
                            "page_no": page_no,
                        }
                    )

            elif item_type == "TableItem":
                md_content = self._table_to_markdown(item, doc)
                table_data = self._table_to_structured(item)
                elements.append(
                    {
                        "type": "table",
                        "summary": "",
                        "content": md_content,
                        "page_no": page_no,
                        "table_data": table_data,
                    }
                )

            elif item_type == "PictureItem":
                image_data = self._extract_picture(item, output_dir, page_no)
                if image_data:
                    elements.append(image_data)

        # Merge header/footer elements at the start
        return hf_elements + elements

    # ── python-docx OOXML helpers ──────────────────────────────

    @staticmethod
    def _extract_ooxml_headers_footers(doc_path: str) -> list[dict[str, Any]]:
        """Extract headers and footers from .docx using python-docx."""
        elements: list[dict[str, Any]] = []
        try:
            from docx import Document as DocxDocument

            docx = DocxDocument(doc_path)

            for section_idx, section in enumerate(docx.sections):
                # Header
                header = section.header
                if header and not header.is_linked_to_previous:
                    text_parts = []
                    for para in header.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text.strip())
                    if text_parts:
                        elements.append(
                            {
                                "type": "header",
                                "content": " | ".join(text_parts),
                                "page_scope": "all",
                                "page_no": section_idx + 1,
                            }
                        )

                # Footer
                footer = section.footer
                if footer and not footer.is_linked_to_previous:
                    text_parts = []
                    for para in footer.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text.strip())
                    if text_parts:
                        elements.append(
                            {
                                "type": "footer",
                                "content": " | ".join(text_parts),
                                "page_scope": "all",
                                "page_no": section_idx + 1,
                            }
                        )
        except Exception:
            pass
        return elements

    @staticmethod
    def _detect_ooxml_page_breaks(doc_path: str) -> set[int]:
        """Detect page break positions in .docx body paragraphs.

        Returns set of paragraph indices (1-based) where page breaks occur.
        This approximates page boundaries — actual page layout depends on
        printer settings and is not stored in OOXML.
        """
        breaks: set[int] = set()
        try:
            from docx import Document as DocxDocument
            from docx.oxml.ns import qn

            docx = DocxDocument(doc_path)
            para_idx = 0
            for para in docx.paragraphs:
                para_idx += 1
                # Check for explicit page break before paragraph
                for run in para.runs:
                    br_elems = run._r.findall(qn("w:br"))
                    for br in br_elems:
                        if br.get(qn("w:type")) == "page":
                            breaks.add(para_idx)
                    # Also check lastRenderedPageBreak
                    for br in run._r.findall(qn("w:lastRenderedPageBreak")):
                        breaks.add(para_idx)
        except Exception:
            pass
        return breaks

    # ── Docling item helpers ────────────────────────────────────

    @staticmethod
    def _get_page_no(item: Any) -> int:
        """Extract page number from a Docling item (1-indexed)."""
        prov = getattr(item, "prov", None)
        if prov and len(prov) > 0:
            page = getattr(prov[0], "page_no", None) or getattr(prov[0], "page", None)
            if page is not None:
                return int(page)
        return 1

    @staticmethod
    def _table_to_markdown(item: Any, doc: Any = None) -> str:
        """Convert a Docling TableItem to markdown string."""
        # Try export_to_markdown with doc argument (required since v2)
        export_fn = getattr(item, "export_to_markdown", None)
        if export_fn:
            try:
                if doc is not None:
                    return str(export_fn(doc=doc))
            except TypeError:
                pass
            try:
                return str(export_fn())
            except Exception:
                pass

        # Fallback: build from grid data
        data = getattr(item, "data", None)
        if not data:
            return getattr(item, "text", "") or ""

        grid = data.grid
        if not grid:
            return ""

        lines = []
        for row_idx, row in enumerate(grid):
            cells = [cell.text for cell in row]
            lines.append("| " + " | ".join(cells) + " |")
            if row_idx == 0:
                lines.append("| " + " | ".join("---" for _ in cells) + " |")

        return "\n".join(lines)

    @staticmethod
    def _table_to_structured(item: Any) -> dict[str, Any]:
        """Extract structured table data from a Docling TableItem."""
        data = getattr(item, "data", None)
        if not data:
            return {"headers": [], "rows": [], "num_rows": 0, "num_cols": 0}

        grid = data.grid
        if not grid:
            return {"headers": [], "rows": [], "num_rows": 0, "num_cols": 0}

        headers = [cell.text for cell in grid[0]]
        rows = [[cell.text for cell in row] for row in grid[1:]]

        return {
            "headers": headers,
            "rows": rows,
            "num_rows": len(rows),
            "num_cols": len(headers),
        }

    @staticmethod
    def _extract_picture(item: Any, output_dir: str, page_no: int) -> dict[str, Any] | None:
        """Extract image data from a Docling PictureItem."""
        image = getattr(item, "image", None)
        if not image:
            return None

        pil_image = getattr(image, "pil_image", None)
        if not pil_image:
            return None

        import io

        buf = io.BytesIO()
        pil_image.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode("utf-8")

        image_id = f"page{page_no:03d}_docling_{id(item)}"
        return {
            "type": "image",
            "image_name": f"{image_id}.png",
            "content": "",
            "image": b64,
            "page_no": page_no,
            "image_id": image_id,
            "bbox": (0.0, 0.0, 0.0, 0.0),
        }
